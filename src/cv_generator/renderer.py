import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION

try:
    from .styles import CVStyles
except ImportError:
    from styles import CVStyles


class DocumentRenderer:
    def __init__(self, photo_path=None):
        self.doc = Document()
        self.styles = CVStyles
        self.photo_path = photo_path
        self._setup_margins()

    def _setup_margins(self):
        section = self.doc.sections[0]
        section.top_margin = self.styles.MARGIN_TOP
        section.bottom_margin = self.styles.MARGIN_BOTTOM
        section.left_margin = self.styles.MARGIN_LEFT
        section.right_margin = self.styles.MARGIN_RIGHT

    def _add_bottom_border(self, paragraph):
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "2A5C5D")
        pBdr.append(bottom)
        paragraph._p.get_or_add_pPr().append(pBdr)

    def add_section_header(self, title, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = self.doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title)
        r.bold = True
        r.font.size = self.styles.SIZE_HEADER
        r.font.name = self.styles.FONT_NAME
        r.font.color.rgb = self.styles.PRIMARY_COLOR
        self._add_bottom_border(p)

    def render_header(self, name, titles, contact_info):
        """Renders the top header with contact info and photo."""
        table = self.doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(6.0)
        table.columns[1].width = Inches(1.5)

        cell_left = table.cell(0, 0)
        cell_right = table.cell(0, 1)

        # Name
        name_para = cell_left.paragraphs[0]
        name_run = name_para.add_run(name)
        name_run.bold = True
        name_run.font.size = self.styles.SIZE_NAME
        name_run.font.name = self.styles.FONT_NAME
        name_run.font.color.rgb = self.styles.PRIMARY_COLOR

        # Titles
        title_para = cell_left.add_paragraph()
        title_run = title_para.add_run(titles)
        title_run.font.size = self.styles.SIZE_TITLE
        title_run.font.name = self.styles.FONT_NAME
        title_run.bold = True

        # Contact
        contact_para = cell_left.add_paragraph()
        contact_run = contact_para.add_run(contact_info)
        contact_run.font.size = self.styles.SIZE_CONTACT
        contact_run.font.name = self.styles.FONT_NAME

        # Photo
        if self.photo_path:
            photo_para = cell_right.paragraphs[0]
            photo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            photo_para.add_run().add_picture(self.photo_path, width=Inches(1.2))

    def render_summary(self, summary_text):
        self.add_section_header("PROFESSIONAL SUMMARY")
        p = self.doc.add_paragraph()
        r = p.add_run(summary_text)
        r.font.name = self.styles.FONT_NAME
        r.font.size = self.styles.SIZE_BODY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(8)

    def setup_columns(self):
        new_section = self.doc.add_section(WD_SECTION.CONTINUOUS)
        cols = OxmlElement("w:cols")
        cols.set(qn("w:num"), "2")
        cols.set(qn("w:equalWidth"), "0")
        col1 = OxmlElement("w:col")
        col1.set(qn("w:w"), "7200")  # 65% width
        col1.set(qn("w:space"), "400")
        col2 = OxmlElement("w:col")
        col2.set(qn("w:w"), "3200")  # 35% width
        cols.append(col1)
        cols.append(col2)
        new_section._sectPr.append(cols)

    def render_experience(self, experience_list):
        self.add_section_header("WORK EXPERIENCE")
        for job in experience_list:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)

            # Company and Role
            r1 = p.add_run(job["organization"])
            r1.bold = True
            r1.font.name = self.styles.FONT_NAME
            r1.font.size = self.styles.SIZE_BOLD_BODY

            r2 = p.add_run(f" – {job['role']}")
            r2.font.name = self.styles.FONT_NAME
            r2.font.size = self.styles.SIZE_BOLD_BODY
            r2.italic = True

            # Dates
            pPr = p._p.get_or_add_pPr()
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "right")
            tab.set(qn("w:pos"), "7200")
            tabs.append(tab)
            pPr.append(tabs)

            r3 = p.add_run(
                f"\t{job.get('start_date', '')} – {job.get('end_date', 'Present')}"
            )
            r3.font.name = self.styles.FONT_NAME
            r3.font.size = self.styles.SIZE_BODY
            r3.italic = True
            r3.font.color.rgb = self.styles.SECONDARY_COLOR

            # Achievements
            if "achievements" in job:
                for achievement in job["achievements"]:
                    bp = self.doc.add_paragraph(achievement, style="List Bullet")
                    bp.paragraph_format.space_after = Pt(2)
                    bp.paragraph_format.left_indent = Inches(0.15)
                    if bp.runs:
                        bp.runs[0].font.name = self.styles.FONT_NAME
                        bp.runs[0].font.size = self.styles.SIZE_BODY

    def render_education(self, education_list):
        self.add_section_header("EDUCATION")
        for edu in education_list:
            p = self.doc.add_paragraph()
            program = edu.get("field") or edu.get("degree", "")
            specialization = edu.get("specialization")
            label = f"{program} ({specialization})" if specialization else program
            r1 = p.add_run(label)
            r1.bold = True
            r1.font.name = self.styles.FONT_NAME
            r1.font.size = self.styles.SIZE_BOLD_BODY
            r2 = p.add_run(
                f" - {edu.get('institution', '')}, {edu.get('location', '')}"
            )
            r2.font.name = self.styles.FONT_NAME
            r2.font.size = self.styles.SIZE_BODY
            dates = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}"
            if dates.strip(" -"):
                d = self.doc.add_paragraph(dates)
                if d.runs:
                    d.runs[0].italic = True
                    d.runs[0].font.size = self.styles.SIZE_SMALL

    def render_publications(self, publications):
        if not publications:
            return
        self.add_section_header("PUBLICATIONS")
        for pub in sorted(publications, key=lambda x: x.get("year", 0), reverse=True):
            venue = pub.get("venue", "")
            title = pub.get("title", "")
            year = pub.get("year", "")
            line = f"{year} - {title}"
            if venue:
                line = f"{line}. {venue}"
            self.doc.add_paragraph(line, style="List Bullet")

    def render_languages(self, languages):
        if not languages:
            return
        self.add_section_header("LANGUAGES")
        for lang in languages:
            name = lang.get("name", "")
            level = lang.get("level", "")
            note = lang.get("note", "")
            text = f"{name}: {level}"
            if note:
                text = f"{text} ({note})"
            self.doc.add_paragraph(text, style="List Bullet")

    def render_skills(self, skills):
        if not skills:
            return
        self.add_section_header("TECHNICAL SKILLS")
        for category, items in skills.items():
            label = category.replace("_", " ").title()
            line = f"{label}: {', '.join(items)}"
            self.doc.add_paragraph(line, style="List Bullet")

    # Add other modular functions here (education, skills, etc) ...

    def save(self, output_path):
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
