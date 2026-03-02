import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION

doc = Document()

# Set narrow margins
section = doc.sections[0]
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

# Add header table for ATS-safe image placement
header_table = doc.add_table(rows=1, cols=2)
header_table.autofit = False
header_table.columns[0].width = Inches(6.0)
header_table.columns[1].width = Inches(1.5)

cell_left = header_table.cell(0, 0)
cell_right = header_table.cell(0, 1)

# Add name
name_para = cell_left.paragraphs[0]
name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
name_run = name_para.add_run('JUAN PABLO RUIZ RODRIGUEZ')
name_run.bold = True
name_run.font.size = Pt(22)
name_run.font.name = 'Arial'
name_run.font.color.rgb = RGBColor(42, 92, 93) # Dark Teal

# Add titles
title_para = cell_left.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
title_run = title_para.add_run('MACHINE LEARNING ENGINEER | RESEARCH ASSISTANT')
title_run.font.size = Pt(11)
title_run.font.name = 'Arial'
title_run.bold = True
title_para.paragraph_format.space_after = Pt(2)

# Add contact info
contact_para = cell_left.add_paragraph()
contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
contact_run = contact_para.add_run('Berlin, Germany | juanpablo.ruiz.r@gmail.com | +49 152 22144630\nlinkedin.com/in/juanpabloruizr | github.com/jpcosec')
contact_run.font.size = Pt(9.5)
contact_run.font.name = 'Arial'
contact_para.paragraph_format.space_after = Pt(12)

# Add photo to right cell (Using circular crop to look more professional)
photo_para = cell_right.paragraphs[0]
photo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
photo_run = photo_para.add_run()
photo_path = os.path.join(os.path.dirname(__file__), '..', 'data', 'reference_data', 'application_assets', 'Photoportrait_circular.png')
photo_run.add_picture(photo_path, width=Inches(1.2))

def add_header(title, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Arial'
    r.font.color.rgb = RGBColor(42, 92, 93)
    
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2A5C5D')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)

# PROFESSIONAL SUMMARY (Full Width)
add_header('PROFESSIONAL SUMMARY')
p = doc.add_paragraph("Machine Learning Engineer and Research Associate bridging applied AI research and production-oriented data systems. Experienced in Python, workflow orchestration (Airflow), and MLOps lifecycle practices (versioning, deployment, monitoring). Background in IoT, ROS, and FAIR-aligned experimentation. Proven experience deploying scalable data and ML pipelines across major cloud providers including Azure, AWS, and Google Cloud (Vertex AI, S3, Spark, BigQuery). Fully compliant with MSCA mobility rules with an Anabin/ZAB validated German Master's equivalency.")
p.style.font.name = 'Arial'
p.style.font.size = Pt(9.5)
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(8)

# Add 2-Column Section Break
new_section = doc.add_section(WD_SECTION.CONTINUOUS)
sectPr = new_section._sectPr
cols = OxmlElement('w:cols')
cols.set(qn('w:num'), '2')
# Total width = 8.5" - 1" = 7.5" = 10800 twips
# Let's do 65% Left (7020), 35% Right (3780)
cols.set(qn('w:equalWidth'), '0')
col1 = OxmlElement('w:col')
col1.set(qn('w:w'), '7200')
col1.set(qn('w:space'), '400') # ~0.27 inch gutter
col2 = OxmlElement('w:col')
col2.set(qn('w:w'), '3200')
cols.append(col1)
cols.append(col2)
sectPr.append(cols)

# ----------------- COLUMN 1 (EXPERIENCE & EDUCATION) -----------------
add_header('WORK EXPERIENCE')
jobs = [
    ('AraraDS', 'ML Consultant', '07.2025 – 09.2025', [
        'Built and deployed Python-based data and ML pipelines with Apache Airflow for production orchestration and scheduling',
        'Integrated LLM-assisted extraction workflows within event-driven document pipelines using Kafka',
        'Applied MLOps practices including model versioning, deployment, and monitoring for iterative experimentation'
    ]),
    ('Universidad de Playa Ancha (UPLA)', 'Research Associate', '11.2024 – Present', [
        'Designed Intelligent Tutoring System (ITS) architecture integrating affective sensing, RAG-backed domain knowledge, and pedagogical decision logic',
        'Developed LAETEC-Vision leveraging PyTorch for real-time cognitive-state classification',
        'Designed reproducible evaluation workflows and structured experiment metadata following FAIR data principles for adaptive interventions'
    ]),
    ('Globalconexus', 'Data Engineer', '04.2024 – 06.2025', [
        'Built and maintained scalable enterprise data pipelines leveraging Azure Data Factory, Databricks, and PySpark',
        'Implemented model deployment monitoring and constructed semantic reporting layers in Power BI'
    ]),
    ('Deloitte', 'Data Scientist', '09.2021 – 03.2023', [
        'Developed ML models and integrated backend/API components for applied analytics use cases'
    ]),
    ('Kwali', 'Machine Learning Engineer', '03.2019 – 08.2021', [
        'Built computer vision-based quality assurance and reporting systems',
        'Designed and deployed real-time image processing pipelines on constrained edge devices including Raspberry Pi and Jetson Nano'
    ]),
    ('DSM Noisemaker', 'Electronics Intern', '12.2017 – 03.2018', [
        'Researched digital audio technologies and developed digital control systems for analog sound processing circuits',
        'Prototyped and tested embedded solutions for audio processing and control interfaces'
    ])
]

for company, role, dates, bullets in jobs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    
    r1 = p.add_run(company)
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(10)
    
    r2 = p.add_run(f' – {role}')
    r2.font.name = 'Arial'
    r2.font.size = Pt(10)
    r2.italic = True
    
    # Adding tab for date, pushing to right edge of column 1 (7200 twips = 5 inches)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '7200') # Width of col1
    tabs.append(tab)
    pPr.append(tabs)
    
    r3 = p.add_run('\t' + dates)
    r3.font.name = 'Arial'
    r3.font.size = Pt(9.5)
    r3.italic = True
    r3.font.color.rgb = RGBColor(100, 100, 100)
    
    for bullet in bullets:
        bp = doc.add_paragraph(bullet, style='List Bullet')
        bp.paragraph_format.space_after = Pt(2)
        bp.paragraph_format.left_indent = Inches(0.15)
        bp.style.font.name = 'Arial'
        bp.style.font.size = Pt(9.5)

# Page Break Protection - keep Education & its content together if it wraps badly
add_header('EDUCATION')
ed_p = doc.add_paragraph()
ed_p.paragraph_format.space_before = Pt(8)
ed_p.paragraph_format.space_after = Pt(2)
# Keep with next
ed_p.paragraph_format.keep_with_next = True

r1 = ed_p.add_run('Universidad de Chile')
r1.bold = True
r1.font.name = 'Arial'
r1.font.size = Pt(10)

pPr = ed_p._p.get_or_add_pPr()
tabs = OxmlElement('w:tabs')
tab = OxmlElement('w:tab')
tab.set(qn('w:val'), 'right')
tab.set(qn('w:pos'), '7200')
tabs.append(tab)
pPr.append(tabs)

r2 = ed_p.add_run('\t2011 – 2019')
r2.font.name = 'Arial'
r2.font.size = Pt(9.5)
r2.italic = True
r2.font.color.rgb = RGBColor(100, 100, 100)

ed_p2 = doc.add_paragraph('Electrical Engineering (Computational Intelligence)')
ed_p2.paragraph_format.keep_with_next = True
ed_p2.runs[0].font.name = 'Arial'
ed_p2.runs[0].font.size = Pt(10)
ed_p2.runs[0].italic = True

bp = doc.add_paragraph("Degree Equivalency: Validated by Anabin/ZAB as equivalent to a German Master's/Diplom degree (>300 ECTS)", style='List Bullet')
bp.paragraph_format.space_after = Pt(2)
bp.paragraph_format.left_indent = Inches(0.15)
bp.style.font.name = 'Arial'
bp.style.font.size = Pt(9.5)

bp2 = doc.add_paragraph('Coursework: Differential equation systems, dynamic modeling, control theory, signal processing, and embedded systems', style='List Bullet')
bp2.paragraph_format.space_after = Pt(2)
bp2.paragraph_format.left_indent = Inches(0.15)
bp2.style.font.name = 'Arial'
bp2.style.font.size = Pt(9.5)

# Insert Page Break logic inside continuous column stream: 
p_col = doc.add_paragraph()
p_col.add_run().add_break(WD_BREAK.COLUMN)

# ----------------- COLUMN 2 (SKILLS & PUBLICATIONS) -----------------

add_header('SKILLS')

skills = [
    ('Programming: ', 'Python, JavaScript, Java, SQL, Embedded C, Verilog'),
    ('Orchestration & DevOps: ', 'Apache Airflow, Prefect, Docker, GIT, FastAPI, Kafka, Azure, AWS, Google Cloud'),
    ('Machine Learning: ', 'PyTorch, TensorFlow, OpenCV, Scikit-learn, LLMs, RAG pipelines, Vertex AI'),
    ('Data Engineering: ', 'Spark, Databricks, PostgreSQL, MongoDB, Neo4J, BigQuery'),
    ('Electronics & Robotics: ', 'Arduino, G-code, ROS, SPICE'),
]
for category, items in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(category)
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(9.5)
    
    # We use a line break here to make it clean in the tight right column
    p.add_run('\n')
    
    r2 = p.add_run(items)
    r2.font.name = 'Arial'
    r2.font.size = Pt(9.5)

add_header('LANGUAGES')
langs = [
    ('Spanish', 'Native (C2)'),
    ('English', 'C1'),
    ('German', 'A2 (Improving)')
]
for lang, level in langs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(lang)
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(9.5)
    
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '3200') # Width of right col
    tabs.append(tab)
    pPr.append(tabs)
    
    r3 = p.add_run('\t' + level)
    r3.font.name = 'Arial'
    r3.font.size = Pt(9.5)
    r3.font.color.rgb = RGBColor(100, 100, 100)
    
add_header('PUBLICATIONS')
pub1 = doc.add_paragraph('Exploration of Incremental Synthetic Non-Morphed Images for Single Morphing Attack Detection', style='List Bullet')
pub1.paragraph_format.space_after = Pt(1)
pub1.paragraph_format.left_indent = Inches(0.15)
pub1.style.font.name = 'Arial'
pub1.style.font.size = Pt(9)
pub1_d = doc.add_paragraph('NeurIPS LXAI Workshop 2025')
pub1_d.paragraph_format.space_after = Pt(6)
pub1_d.paragraph_format.left_indent = Inches(0.15)
pub1_d.runs[0].font.name = 'Arial'
pub1_d.runs[0].font.size = Pt(9)
pub1_d.runs[0].italic = True
pub1_d.runs[0].font.color.rgb = RGBColor(100, 100, 100)

pub2 = doc.add_paragraph('Intelligent Tutor for Tasks in the Domain of Linear Algebra', style='List Bullet')
pub2.paragraph_format.space_after = Pt(1)
pub2.paragraph_format.left_indent = Inches(0.15)
pub2.style.font.name = 'Arial'
pub2.style.font.size = Pt(9)
pub2_d = doc.add_paragraph('ICITED 2025')
pub2_d.paragraph_format.space_after = Pt(2)
pub2_d.paragraph_format.left_indent = Inches(0.15)
pub2_d.runs[0].font.name = 'Arial'
pub2_d.runs[0].font.size = Pt(9)
pub2_d.runs[0].italic = True
pub2_d.runs[0].font.color.rgb = RGBColor(100, 100, 100)

output_dir = os.path.join(os.path.dirname(__file__), '..', 'data', 'pipelined_data', 'tu_berlin', '201084', 'cv', 'rendered')
os.makedirs(output_dir, exist_ok=True)
doc.save(os.path.join(output_dir, 'CV_TU_Berlin_201084_Col_ATS.docx'))
