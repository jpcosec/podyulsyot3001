"""DOCX CV builder using python-docx."""

from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from src.core.tools.render.styles import resolve_docx_style

#TODO: Pydocs, would be a good idea being able to use templates here 

class DocumentRenderer: 
    def __init__(
        self,
        photo_path=None,
        template_name: str = "classic",
        template_path: Path | None = None,
    ):
        self.doc = Document(str(template_path)) if template_path else Document()
        self.styles = resolve_docx_style(template_name)
        self.photo_path = photo_path
        self._setup_margins()

    def _setup_margins(self):
        section = self.doc.sections[0]
        section.top_margin = self.styles.MARGIN_TOP
        section.bottom_margin = self.styles.MARGIN_BOTTOM
        section.left_margin = self.styles.MARGIN_LEFT
        section.right_margin = self.styles.MARGIN_RIGHT

    def _make_floating_photo(self, para, photo_path, width=Inches(1.2)):
        """Attach photo as a floating anchored image (top-right of margin).

        The image is inserted into `para`'s run, then converted from inline
        (wp:inline) to anchored (wp:anchor) so it floats visually at top-right
        while all text paragraphs remain in the normal paragraph flow — fully
        readable by ATS.
        """
        run = para.add_run()
        run.add_picture(photo_path, width=width)

        drawing = para._p.find(".//" + qn("w:drawing"))
        if drawing is None:
            return
        inline = drawing.find(qn("wp:inline"))
        if inline is None:
            return

        extent = inline.find(qn("wp:extent"))
        cx, cy = extent.get("cx"), extent.get("cy")

        anchor = OxmlElement("wp:anchor")
        for attr, val in [
            ("distT", "0"),
            ("distB", "0"),
            ("distL", "114300"),
            ("distR", "114300"),
            ("simplePos", "0"),
            ("relativeHeight", "251658240"),
            ("behindDoc", "0"),
            ("locked", "0"),
            ("layoutInCell", "1"),
            ("allowOverlap", "0"),
        ]:
            anchor.set(attr, val)

        sp = OxmlElement("wp:simplePos")
        sp.set("x", "0")
        sp.set("y", "0")
        anchor.append(sp)

        ph = OxmlElement("wp:positionH")
        ph.set("relativeFrom", "margin")
        al = OxmlElement("wp:align")
        al.text = "right"
        ph.append(al)
        anchor.append(ph)

        pv = OxmlElement("wp:positionV")
        pv.set("relativeFrom", "margin")
        pvo = OxmlElement("wp:posOffset")
        pvo.text = "0"
        pv.append(pvo)
        anchor.append(pv)

        ext = OxmlElement("wp:extent")
        ext.set("cx", cx)
        ext.set("cy", cy)
        anchor.append(ext)

        eff = OxmlElement("wp:effectExtent")
        for side in ("l", "t", "r", "b"):
            eff.set(side, "0")
        anchor.append(eff)

        wrap = OxmlElement("wp:wrapSquare")
        wrap.set("wrapText", "bothSides")
        anchor.append(wrap)

        for child_tag in (qn("wp:docPr"), qn("wp:cNvGraphicFramePr")):
            child = inline.find(child_tag)
            if child is not None:
                anchor.append(copy.deepcopy(child))

        graphic_tag = "{http://schemas.openxmlformats.org/drawingml/2006/main}graphic"
        graphic = inline.find(graphic_tag)
        if graphic is not None:
            anchor.append(copy.deepcopy(graphic))

        drawing.remove(inline)
        drawing.append(anchor)

    def _bullet_paragraph(self, text: str, font_size=None) -> None:
        """Add a paragraph prefixed with a standard Unicode bullet (•).

        Using a literal • character instead of Word's List Bullet style
        ensures the bullet is preserved in PDF text extraction as U+2022
        rather than the Wingdings private-use character \\uf0b7.
        """
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        run = p.add_run("• " + text)
        run.font.name = self.styles.FONT_NAME
        run.font.size = font_size or self.styles.SIZE_BODY

    def _add_bottom_border(self, paragraph):
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), self.styles.PRIMARY_COLOR_HEX)
        pBdr.append(bottom)
        paragraph._p.get_or_add_pPr().append(pBdr)

    def add_section_header(self, title, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = self.doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title)
        r.bold = True
        r.font.size = self.styles.SIZE_HEADER
        r.font.name = self.styles.FONT_NAME
        r.font.color.rgb = self.styles.PRIMARY_COLOR
        self._add_bottom_border(p)

    def render_header(self, name, titles, contact_info):
        """Render ATS-safe single-column header with optional floating photo."""
        name_para = self.doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_para.paragraph_format.space_after = Pt(2)
        name_run = name_para.add_run(name)
        name_run.bold = True
        name_run.font.size = self.styles.SIZE_NAME
        name_run.font.name = self.styles.FONT_NAME
        name_run.font.color.rgb = self.styles.PRIMARY_COLOR

        if self.photo_path:
            self._make_floating_photo(name_para, self.photo_path)

        title_para = self.doc.add_paragraph()
        title_para.paragraph_format.space_after = Pt(2)
        title_run = title_para.add_run(titles)
        title_run.font.size = self.styles.SIZE_TITLE
        title_run.font.name = self.styles.FONT_NAME
        title_run.bold = True
        title_run.font.color.rgb = self.styles.SECONDARY_COLOR

        contact_para = self.doc.add_paragraph()
        contact_para.paragraph_format.space_after = Pt(6)
        contact_run = contact_para.add_run(contact_info)
        contact_run.font.size = self.styles.SIZE_CONTACT
        contact_run.font.name = self.styles.FONT_NAME
        contact_run.font.color.rgb = self.styles.SECONDARY_COLOR

        self._add_bottom_border(contact_para)

    def render_summary(self, summary_text):
        self.add_section_header("SUMMARY")
        p = self.doc.add_paragraph()
        r = p.add_run(summary_text)
        r.font.name = self.styles.FONT_NAME
        r.font.size = self.styles.SIZE_BODY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(8)

    def render_experience(self, experience_list):
        self.add_section_header("EXPERIENCE")
        for job in experience_list:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)

            r1 = p.add_run(job["organization"])
            r1.bold = True
            r1.font.name = self.styles.FONT_NAME
            r1.font.size = self.styles.SIZE_BOLD_BODY

            r2 = p.add_run(f" – {job['role']}")
            r2.font.name = self.styles.FONT_NAME
            r2.font.size = self.styles.SIZE_BOLD_BODY
            r2.italic = True

            pPr = p._p.get_or_add_pPr()
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "right")
            tab.set(qn("w:pos"), "7200")
            tabs.append(tab)
            pPr.append(tabs)

            r3 = p.add_run(
                f"\t{job.get('start_date', '')} – {job.get('end_date', 'Present')}"
            )
            r3.font.name = self.styles.FONT_NAME
            r3.font.size = self.styles.SIZE_BODY
            r3.italic = True
            r3.font.color.rgb = self.styles.SECONDARY_COLOR

            if "achievements" in job:
                for achievement in job["achievements"]:
                    self._bullet_paragraph(achievement)

    def render_education(self, education_list):
        self.add_section_header("EDUCATION")
        for edu in education_list:
            p = self.doc.add_paragraph()
            program = edu.get("field") or edu.get("degree", "")
            specialization = edu.get("specialization")
            label = f"{program} ({specialization})" if specialization else program
            r1 = p.add_run(label)
            r1.bold = True
            r1.font.name = self.styles.FONT_NAME
            r1.font.size = self.styles.SIZE_BOLD_BODY
            r2 = p.add_run(
                f" - {edu.get('institution', '')}, {edu.get('location', '')}"
            )
            r2.font.name = self.styles.FONT_NAME
            r2.font.size = self.styles.SIZE_BODY
            dates = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}"
            if dates.strip(" -"):
                d = self.doc.add_paragraph(dates)
                if d.runs:
                    d.runs[0].italic = True
                    d.runs[0].font.size = self.styles.SIZE_SMALL

            equivalency_note = edu.get("equivalency_note", "")
            if equivalency_note:
                note = self.doc.add_paragraph()
                note.paragraph_format.space_after = Pt(2)
                note_run = note.add_run(equivalency_note)
                note_run.font.name = self.styles.FONT_NAME
                note_run.font.size = self.styles.SIZE_SMALL

    def render_publications(self, publications):
        if not publications:
            return
        self.add_section_header("PUBLICATIONS")
        for pub in sorted(publications, key=lambda x: x.get("year", 0), reverse=True):
            venue = pub.get("venue", "")
            title = pub.get("title", "")
            year = pub.get("year", "")
            line = f"{year} - {title}"
            if venue:
                line = f"{line}. {venue}"
            self._bullet_paragraph(line)

    def render_languages(self, languages):
        if not languages:
            return
        self.add_section_header("LANGUAGES")
        for lang in languages:
            name = lang.get("name", "")
            level = lang.get("level", "")
            note = lang.get("note", "")
            text = f"{name}: {level}"
            if note:
                text = f"{text} ({note})"
            self._bullet_paragraph(text, font_size=self.styles.SIZE_BODY)

    def render_skills(self, skills):
        if not skills:
            return
        self.add_section_header("SKILLS")
        for category, items in skills.items():
            label = category.replace("_", " ").title()
            line = f"{label}: {', '.join(items)}"
            self._bullet_paragraph(line)

    def save(self, output_path):
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
