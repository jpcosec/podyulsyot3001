"""Motivation letter renderer: markdown → DOCX → PDF via LibreOffice."""

from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


def _add_runs(paragraph, text: str, base_size: Pt, font_name: str = "Calibri") -> None:
    """Add runs to a paragraph, handling **bold** markdown inline markers."""
    parts = _BOLD_PATTERN.split(text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = font_name
        run.font.size = base_size
        run.bold = (i % 2 == 1)  # odd indices are inside **...**


def _parse_letter(md_text: str) -> dict:
    """Extract structured fields from the motivation letter markdown."""
    lines = md_text.splitlines()

    # Find the two --- separators
    sep_indices = [i for i, ln in enumerate(lines) if ln.strip() == "---"]

    # Everything between first and second separator is the header block
    header_block = lines[sep_indices[0] + 1 : sep_indices[1]]
    date = ""
    for ln in header_block:
        stripped = ln.strip()
        if stripped and not stripped.startswith("**"):
            date = stripped

    # Everything after second separator is the letter body
    body_lines = lines[sep_indices[1] + 1 :]

    # Split body into paragraphs on blank lines
    paragraphs = []
    current: list[str] = []
    for ln in body_lines:
        if ln.strip() == "":
            if current:
                paragraphs.append(" ".join(current))
                current = []
        else:
            current.append(ln.strip())
    if current:
        paragraphs.append(" ".join(current))

    header_lines: list[str] = []
    if paragraphs and "dear " not in paragraphs[0].lower():
        header_lines = [line.strip() for line in body_lines[:5] if line.strip()]
        paragraphs = paragraphs[1:]

    # Find closing: "Yours sincerely," and everything after
    closing_index = next(
        (i for i, p in enumerate(paragraphs) if "sincerely" in p.lower()), -1
    )
    body_paragraphs = paragraphs[:closing_index] if closing_index != -1 else paragraphs
    closing_block = paragraphs[closing_index:] if closing_index != -1 else []

    return {
        "date": date,
        "header_lines": header_lines,
        "paragraphs": body_paragraphs,
        "closing_block": closing_block,
    }


def render_letter_docx(md_path: str | Path, output_path: str | Path) -> Path:
    """Render motivation letter markdown to DOCX."""
    md_text = Path(md_path).read_text(encoding="utf-8")
    fields = _parse_letter(md_text)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    font = "Calibri"
    size_body = Pt(11)
    size_small = Pt(10)
    gray = RGBColor(80, 80, 80)

    def _para(text="", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            _add_runs(p, text, size_body, font)
        return p

    # Date (right-aligned)
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dp.paragraph_format.space_after = Pt(12)
    r = dp.add_run(fields["date"])
    r.font.name = font
    r.font.size = size_body
    r.font.color.rgb = gray

    for index, line in enumerate(fields.get("header_lines", [])):
        contact = doc.add_paragraph()
        contact.paragraph_format.space_after = Pt(0 if index < len(fields["header_lines"]) - 1 else 12)
        run = contact.add_run(line)
        run.font.name = font
        run.font.size = size_body

    # Horizontal rule via border
    hr = doc.add_paragraph()
    hr.paragraph_format.space_after = Pt(12)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    hr._p.get_or_add_pPr().append(pBdr)

    # Body paragraphs
    for para_text in fields["paragraphs"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        _add_runs(p, para_text, size_body, font)

    # Closing block
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    for line in fields["closing_block"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _add_runs(p, line, size_small, font)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def convert_docx_to_pdf(docx_path: str | Path, output_dir: str | Path) -> Path:
    """Convert DOCX to PDF using LibreOffice headless."""
    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir",
         str(output_dir), str(docx_path)],
        capture_output=True,
        text=True,
    )
    pdf_path = Path(output_dir) / (Path(docx_path).stem + ".pdf")
    if not pdf_path.exists():
        raise RuntimeError(
            f"LibreOffice conversion failed:\n{result.stderr}"
        )
    return pdf_path
